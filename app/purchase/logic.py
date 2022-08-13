import os
from math import ceil

from docx import Document
from flask import redirect, url_for, flash
from flask_babel import _
from flask_login import current_user
from sqlalchemy import or_, not_

from app import session
from app.containers import ConsumptionItemState
from app.festival.logic import load_participants_from_db
from app.main.utils import send_file
from app.models import PackagingUnitType, ConsumptionItem, Festival, User
from config import Config


def get_pku_selection():
    result = [(-1, "")]
    types = session.query(PackagingUnitType).all()
    for t in types:
        result.append((t.id, t.abbreviation))
    return result


def get_festivals():
    result = [(-1, "")]
    festivals = session.query(Festival).filter(not_(Festival.is_closed)).all()
    for f in festivals:
        result.append((f.id, f.title))
    return result


def calculate_redirect(item):
    if item.state == ConsumptionItemState.stock:
        return redirect(url_for("purchase.stock_overview"))
    if item.state == ConsumptionItemState.wishlist:
        return redirect(url_for("purchase.wishlist"))
    return redirect(url_for("main.index"))


def check_shopping_empty():
    shopping_list_len = session.query(ConsumptionItem).filter(or_(
        ConsumptionItem.state == ConsumptionItemState.purchase,
        ConsumptionItem.state == ConsumptionItemState.cart
    )).count()
    return shopping_list_len == 0


def get_pallets(amount, size, days):
    return ceil((amount * days) / size)


def calculate_drinks(festival_id, is_testing):
    festival = session.query(Festival).get(festival_id)
    duration = festival.end_date - festival.start_date
    days = duration.days
    beer_amount = mixed_amount = water_amount = 0
    for p in load_participants_from_db(festival_id):
        beer_amount += p.beer_demand
        mixed_amount += p.mixed_demand
        water_amount += p.water_demand
    beer_large_pallets = get_pallets(beer_amount, 24, days)
    mixed_large_pallets = get_pallets(mixed_amount, 24, days)
    beer_small_pallets = get_pallets(beer_amount, 18, days)
    mixed_small_pallets = get_pallets(mixed_amount, 18, days)
    water_pallets = get_pallets(water_amount, (6 * 1.5), days)

    six = session.query(PackagingUnitType).filter_by(internal_name="Sixpacks").first()
    cns = session.query(PackagingUnitType).filter_by(internal_name="Cans").first()

    beer_info = "{}x24 or {}x18".format(beer_large_pallets, beer_small_pallets)
    mixed_info = "{}x24 or {}x18".format(mixed_large_pallets,
                                         mixed_small_pallets)
    water_info = "6x1.5l"

    user = current_user
    if is_testing:
        user = session.query(User).get(1)

    beer = ConsumptionItem(name=_("Beer"), pku_id=cns.id,
                           amount=(beer_amount * days),
                           info=beer_info,
                           requestor=user)
    mixed = ConsumptionItem(name=_("Mixed"), pku_id=cns.id,
                            amount=(mixed_amount * days),
                            info=mixed_info,
                            requestor=user)
    water = ConsumptionItem(name=_("Water"), pku_id=six.id,
                            amount=water_pallets,
                            info=water_info,
                            requestor=user)
    session.add_all([beer, mixed, water])
    session.commit()


def generate_shopping_list(festival_id, is_testing=False):
    calculate_drinks(festival_id, is_testing)
    requested_items = session.query(ConsumptionItem).filter_by(
        state=ConsumptionItemState.wishlist).all()
    for r in requested_items:
        r.festival_id = festival_id
        available: ConsumptionItem = session.query(ConsumptionItem).filter(
            ConsumptionItem.state == ConsumptionItemState.stock,
            ConsumptionItem.name == r.name
        ).first()
        if available:
            adjust_amount(available, r)
        else:
            r.state = ConsumptionItemState.purchase
    session.commit()
    if not is_testing:
        flash(_("List generated."))
        return redirect(url_for("purchase.shopping_list", festival_id=festival_id))


def adjust_amount(available: ConsumptionItem, requested: ConsumptionItem):
    p_amount = requested.amount - available.amount
    if p_amount <= 0:
        # demand can be satisfied from stock
        session.delete(requested)
    else:
        requested.amount = p_amount
        requested.state = ConsumptionItemState.purchase


def export_and_download_docx():
    lines = session.query(ConsumptionItem, PackagingUnitType) \
        .join(PackagingUnitType).filter(
        ConsumptionItem.state == ConsumptionItemState.purchase).all()
    length = len(lines)
    if length == 0:
        return

    festival_id: int = lines[0].ConsumptionItem.festival_id
    festival = session.query(Festival).filter_by(id=festival_id).first()

    document = Document()
    document.add_heading(festival.title)
    rows = length + 1
    columns = 3
    table = document.add_table(rows=rows, cols=columns)
    table.style = "LightShading-Accent1"
    table_headers = [_("Article"), _("Amount"), _("Info")]
    for r in range(rows):
        for c in range(columns):
            cell = table.cell(r, c)
            if r == 0:
                cell.text = table_headers[c]
            else:
                ci: ConsumptionItem = lines[r - 1].ConsumptionItem
                if c == 0:
                    cell.text = ci.name
                elif c == 1:
                    pku: PackagingUnitType = lines[r - 1].PackagingUnitType
                    cell.text = "{} {}".format(ci.amount, pku.abbreviation)
                else:
                    if ci.info:
                        cell.text = ci.info

    filename = "{}.docx".format(festival.title)
    file_path = os.path.join(Config.SHOPPING_LIST_PATH, filename)
    document.save(file_path)

    return send_file(file_path, filename, "text/docx")
